"""生成《个性化挽留策略 A/B 实验设计 v1.0》Word 文档（农场分支已并入风控）。"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parents[2] / "prd" / "ab_testing" / "挽留策略AB实验设计_v1.0.docx"
CN = "PingFang SC"
ACCENT = RGBColor(0x09, 0x69, 0xDA)
MUTED = RGBColor(0x65, 0x6D, 0x76)

doc = Document()

# 默认字体（中西文）
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN)

def _set_cn(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), CN)
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn(run, size=17 if level == 1 else 13.5, bold=True,
            color=None if level == 1 else ACCENT)
    p.space_before = Pt(10)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "1F2328")
        pbdr.append(bottom); pPr.append(pbdr)
    return p

def para(text, size=10.5, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn(run, size=size, bold=bold, color=color)
    run.font.italic = italic
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        _set_cn(run, size=10.5)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        _set_cn(run, size=10, bold=True)
        shade(c, "F0F3F6")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_cn(run, size=9.5)
    doc.add_paragraph()
    return t

def callout(text, fill="FFF8C5"):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; c.text = ""
    shade(c, fill)
    for i, line in enumerate(text.split("\n")):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        run = p.add_run(line); _set_cn(run, size=9.5)
    doc.add_paragraph()

# ---------- 封面 / 元信息 ----------
title = para("个性化挽留策略 A/B 实验设计（HMM 状态分层）· v1.0", size=18, bold=True)
pPr = title._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "6"); bottom.set(qn("w:color"), "1F2328")
pbdr.append(bottom); pPr.append(pbdr)
for m in ["游戏：FM01（捕鱼大师） · 数据表：transform-agfish-game.public.bullet",
          "配套：A/B 实验平台 PRD v0.2 · 分组方案 AB分组方案设计_v0.1.md",
          "起草：2026-08-13 · 状态：draft（时间线为示例，正式日期以排期为准）",
          "layer：player_value_intervention · mutex_group：retention_intervention"]:
    para(m, size=9, color=MUTED)

# ---------- 1 假设 ----------
heading("1. 假设（Hypothesis）")
callout("给处于不同生命周期状态（HMM：T1 / S1-low / S2-engaged / S3-escaped）的玩家开启个性化挽留策略"
        "（ON），相较不开启（OFF/holdout），能显著提升其 D7 留存与转化，且不恶化净收入、付费与风控护栏。",
        fill="EAF2FB")
para("派生子假设（每 HMM 状态独立检验）：", bold=True)
bullets([
    "H-S3（escaped/流失）：挽留对高流失风险人群提升最大（重点验证）。",
    "H-S1（low）：挽留能把低活跃拉回活跃。",
    "H-S2（engaged）：挽留对已活跃人群边际提升有限（可能无效，验证是否浪费成本）。",
    "H-T1（初始）：挽留对新客/首触早期留存的影响（含义待确认）。",
])
para("启动后冻结，改变假设 → 新建实验。", italic=True, color=MUTED)

# ---------- 2 背景 ----------
heading("2. 背景与目标")
para("个性化挽留系统已上线并按生命周期给玩家打 CR_FISHING_V1:Tk 标签、施加加成，但当前没有干净对照来量化"
     "“挽留带来多少增量、对哪类人有效”。本实验建立挽留 ON vs OFF 的随机对照，并用处理无关的 HMM 状态分层，"
     "回答：挽留的因果增量，以及对哪种生命周期状态最该投入。")

# ---------- 3 人群 ----------
heading("3. 分析人群与分组")
bullets([
    "subject_type：FM01 玩家账号（当前用户 ID 体系，不依赖 OneID）。",
    "eligibility（筛选逻辑）：① 排除 风控（sticky，套利农场已含在内）；② 排除 dynamic_rtp 臂（吃 RTP，混淆）；"
    "③ 剩余 = 挽留臂(ON) + holdout(OFF)，两臂均 RTP-off。",
    "分层维度（HMM，处理无关，冻结）：T1 / S1 low / S2 engaged / S3 escaped，从行为序列算出、与是否被挽留无关"
    " → 两臂都能算 → 无 post-treatment 偏差。",
    "处理臂：Control = 挽留 OFF（holdout）；Variant = 挽留策略 ON。",
    "终端格：4 状态 × 2 臂 = 8 格。",
])
para("分割逻辑：全体 →（风控? 是→出局 / 否）→ 实验人群 →（dynamic_rtp 臂：单独 RTP 实验）+"
     "（挽留 AB：挽留ON+holdout OFF，均 RTP-off）→ 按 HMM 状态 T1/S1/S2/S3 各分 OFF|ON。", color=MUTED, size=9.5)

# ---------- 4 变量 ----------
heading("4. 变量（Control / Variant）")
table(["", "内容", "约束"], [
    ["Control", "不施加个性化挽留（holdout，保持基准体验）", "可追溯的基准配置版本"],
    ["Variant", "开启现行个性化挽留策略 CR_FISHING_V1（含加成/触达）", "Phase 1 仅一个 Variant；策略实质变化则结束重开"],
])
para("若后续比“挽留策略 A vs B”，属另一实验（两臂皆 ON，在挽留臂内再切），本实验不含。", italic=True, color=MUTED, size=9.5)

# ---------- 5 指标 ----------
heading("5. 指标")
para("5.1 主指标（启动前冻结）", bold=True, color=ACCENT)
para("D7 留存率：cohort 玩家在入组日 +7 自然日仍有任意 FM01 有效投注的比例。统计主体=玩家；分析人群=ITT；"
     "窗口=D7；时区=北京(UTC+8)。")
para("5.2 次要指标（观察，不判胜负）", bold=True, color=ACCENT)
para("D1、D3 留存；状态跃迁率（S3→S2/S1 回流、S1→S2 提升）；人均投注次数/额。")
para("5.3 护栏指标（每项设停止阈值）", bold=True, color=ACCENT)
table(["护栏", "口径", "停止阈值（示例，待批准）"], [
    ["人均净收入 / ARPU", "(投注−派彩)/人", "Variant 相对 Control 下降 > 5%"],
    ["付费率", "付费用户占比", "下降 > 3pp"],
    ["整层 RTP", "派彩/投注", "Variant RTP 越界（如 >100%，防放水被套利）"],
    ["客诉 / 投诉率", "投诉工单/人", "显著上升"],
    ["曝光完整性", "exposure/assignment", "< 阈值触发数据质量 invalid"],
])

# ---------- 6 统计 ----------
heading("6. 统计设计")
bullets([
    "analysis_population：ITT（默认）。未曝光的 assignment 仍进主分析，曝光完整性作诊断。",
    "固定周期：预设 minimum_duration 与目标样本，到达前不因普通显著性提前停。",
    "alpha = 0.05（双侧），power = 0.8（具体值由数据负责人批准）。",
    "分层分析：每个 HMM 状态内做两比例检验（ON vs OFF）；多状态用预注册多重比较校正（如 Holm）。",
])
para("6.1 每臂样本量需求（主指标 D7 留存）", bold=True, color=ACCENT)
para("n/arm = (z_α/2 + z_β)² · [p1(1-p1)+p2(1-p2)] / (p1-p2)²", size=9.5)
table(["HMM 状态", "基线 D7（待标定）", "MDE +3pp", "MDE +5pp", "MDE +8pp"], [
    ["T1 初始", "15%", "2,399", "903", "374"],
    ["S1 low", "10%", "1,771", "683", "291"],
    ["S2 engaged", "30%", "3,760", "1,374", "546"],
    ["S3 escaped", "5%", "1,056", "432", "197"],
])
callout("关键约束：需求为每臂每状态。Control 来自 holdout（人群较小），是样本瓶颈——单波不够则跨波累计到目标样本再读数。"
        "最终数字待 HMM 各状态实际人数分布 + 真实基线 D7 标定后锁定。脚本：jobs/ab_testing/ab_sample_size.py。")

# ---------- 7 数据质量 ----------
heading("7. 数据质量与 SRM")
bullets([
    "双层 SRM：assignment 层 + exposure 层分别比对实际 vs 预设比例。",
    "分状态 SRM：8 个终端格实际比例与预期比对，任一 P0 失败 → 结论锁定。",
    "曝光口径：ON/OFF 两臂同一曝光触发口径，去重。",
    "右截断：末尾入组 cohort 的 D7 未成熟时不计入主分析（见 §9）。",
])

# ---------- 8 灰度 ----------
heading("8. 灰度扩量与紧急停止")
bullets([
    "ramp_plan：小流量起步 → 每阶段检查（SRM/护栏/RTP/QA/上阶段最短观察）→ 逐步扩量；扩量不改变已分组用户组别。",
    "emergency_stop_rule：触发 = 护栏越界 / 严重 SRM / RTP 异常放水 / 客诉激增。停止后新增决策回到基准（挽留 OFF）；"
    "已产生加成/订单不回滚；恢复需新批准。",
])

# ---------- 9 时间线 ----------
heading("9. 时间线与验证窗口（示例，日期以排期为准）")
table(["阶段", "日期（示例）", "内容 / 验证"], [
    ["G0 设计冻结", "2026-08-20", "假设/指标/HMM 口径/护栏阈值/样本量关闭 to_be_confirmed"],
    ["QA + 数据链路就绪", "08-20 ~ 08-27", "测试身份验证分组；assignment/exposure 双事件打通"],
    ["灰度 ramp 起步", "2026-08-28", "小流量；查 SRM/护栏/RTP，健康后扩量"],
    ["全量运行", "2026-09-04", "达比例上限；进入正式 cohort 累计"],
    ["最短运行 minimum_duration", "14 天", "09-04 ~ 09-18 主 cohort 入组窗口（覆盖 2 个完整周周期）"],
    ["D7 验证窗口成熟", "2026-09-25", "最后入组 cohort（09-18）+7 天 → 主指标可读"],
    ["结论评审 readout", "2026-09-26", "主指标/护栏/SRM/区间/样本成熟度出报告"],
    ["决策 & rollout", "2026-09-28+", "ship→灰度全量 / extend / stop / investigate"],
])
callout("验证时间（观测窗口）口径：\n"
        "• 每个 cohort 的 D1/D3/D7 分别在其入组日 +1/+3/+7 成熟。\n"
        "• 主指标 D7 需等窗口内最后一个 cohort 也满 7 天才完整（示例 09-25）。\n"
        "• 实验末尾 7 天入组的 cohort，D7 未成熟 → 右截断，不进主分析或触发 extend。\n"
        "• minimum_duration 取 14 天以覆盖完整周内节律（工作日/周末差异），避免周期偏差。")

# ---------- 10 决策 ----------
heading("10. 决策框架（ship / extend / stop / investigate）")
table(["主指标(D7)", "护栏", "数据质量", "结论"], [
    ["达统计+业务门槛", "可接受", "正常", "ship（该状态开挽留）"],
    ["方向对但样本/窗口不足", "—", "正常", "extend"],
    ["效果≈0 或负向", "—", "正常", "stop（该状态不投挽留，省成本）"],
    ["正向", "越界", "—", "investigate（权衡）"],
    ["—", "—", "无效/联合SRM", "investigate（锁定排查）"],
])
para("允许分状态差异化结论：如 S3/S1 ship、S2 stop（对已活跃人群挽留无增量则不投）。", italic=True, color=MUTED, size=9.5)

# ---------- 11 风险 ----------
heading("11. 风险与开放问题")
table(["编号", "问题 / 风险", "处理"], [
    ["OQ-1", "T1 状态确切含义与 HMM 四状态判定阈值", "需给定义，写死分层口径"],
    ["OQ-2", "HMM 各状态实际人数分布", "标定后回填 §6.1"],
    ["OQ-3", "真实基线 D7（各状态）", "历史数据标定，替换示例值"],
    ["OQ-4", "holdout(OFF) 人群偏小 → 跨波累计或调 MDE", "数据负责人定"],
    ["OQ-5", "挽留 OFF 是否涉及玩家权益/合规（停发已承诺加成）", "风控/合规确认"],
])

# ---------- 12 附录 ----------
heading("12. 附录：数据口径与复算")
bullets([
    "组别/分割口径：AB分组方案设计_v0.1.md",
    "样本量脚本：jobs/ab_testing/ab_sample_size.py",
    "留存/行为口径（时区 UTC+8、BASE_FILTER、右截断）：沿用 jobs/fishing/fm01_grouping.py 及既有留存分析脚本。",
    "HMM 状态：由挽留侧模型产出 T1 / S1 / S2 / S3，作为外部输入接入分层（口径待接口确认）。",
])

para("FM01 个性化挽留策略 A/B 实验设计 v1.0 · 2026-08-13 · draft", size=8.5, color=MUTED)

doc.save(str(OUT))
print("saved:", OUT)
